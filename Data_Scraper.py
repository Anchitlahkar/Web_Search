from os import replace
from typing import final
from selenium import webdriver
from bs4 import BeautifulSoup
import time
import csv
from docx import Document

query = input('Key Word: ')
file_name = input('File Name: ')

start_url = "http://google.com/search?q="+query
browser = webdriver.Chrome(
    "D:/System install files/chromedriver_win32/chromedriver")

time.sleep(3)

browser.get(start_url)

time.sleep(5)

data = []

Soup = BeautifulSoup(browser.page_source, "html.parser")

raw_text = []
updated_text = []


mydoc = Document()


def scrape():

    print('\n\nPages To Search: ', len(
        Soup.find_all("div", attrs={"class", "tF2Cxc"})))

    for div_f_tag in Soup.find_all("div", attrs={"class", "tF2Cxc"}):
        div_S_tags = div_f_tag.find_all("div", attrs={"class", "yuRUbf"})

        temp_list = []
        for index, a_tag in enumerate(div_S_tags):
            if index == 0:
                temp_list.append(a_tag.find_all("a", href=True)[0]["href"])

            else:
                try:
                    temp_list.append(a_tag.contents[0])

                except:
                    temp_list.append('')

        data.append(temp_list)

    for i in range(0, len(data)):
        browser.get(data[i][0])

        time.sleep(5)

        print('\nPage no: ',i+1)

        p =len(BeautifulSoup(
            browser.page_source, "html.parser").find_all("p"))        
        print('p: ', p)

        if p != 0:
            time.sleep(2)
    
            raw_text.append(BeautifulSoup(
                    browser.page_source, "html.parser").find_all("p"))
    
            raw_text.append('...............................................................................................................................................................................................................'+
                                '........................Url:'+str(data[i])+'.....................................................................................................................................................................'+
                            '...............................................................................................................................................................................................................')

        else:
            time.sleep(2)
    
            raw_text.append(BeautifulSoup(
                    browser.page_source, "html.parser").find_all("p"))
    
            raw_text.append('...............................................................................................................................................................................................................'+
                                '................Unable To Featch Results---Please Check The URL,     url: '+str(data[i])+'      .........................................................................................................................................'+
                            '...............................................................................................................................................................................................................')


    with open("url.csv", "w") as f:
        csv_Writter = csv.writer(f)
        csv_Writter.writerows(data)


def remove_tags(text):

    soup = BeautifulSoup(text, "html.parser")

    a_tag_words = []
    temp_text = []

    for data in soup(['img']):
        data.decompose()
        temp_text.append(' '.join(soup.stripped_strings))


    for data in soup("a"):
        a_tag_words.append(data)
        temp_text.append(' '.join(soup.stripped_strings))

    print('\n\n\n', a_tag_words)


    para = temp_text

    try:
        mydoc.add_heading(query)
        mydoc.add_paragraph(para[0])
        mydoc.save(file_name+" Data.docx")

    except:
        mydoc.add_heading(query)
        mydoc.add_paragraph(para)
        mydoc.save(file_name+" Data_1.docx")

    # print('\n\n\n para: ', para)


scrape()
remove_tags(str(raw_text))
print("\n\n\n\n\n\n")
input('exit...')
exit()
