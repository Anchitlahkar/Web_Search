from selenium import webdriver
from webdriver_manager.microsoft import EdgeChromiumDriverManager
from bs4 import BeautifulSoup
import time
import csv
import requests
from docx import Document
import os

def resource_path(relative_path):
    try:
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.dirname(__file__)
    return os.path.join(base_path, relative_path)

query = input('Key Word: ')
file_name = input('File Name: ')


start_url = "https://www.bing.com/"
browser = webdriver.Edge(EdgeChromiumDriverManager().install())


time.sleep(5)

browser.get(start_url)

time.sleep(5)

browser.get(start_url+"search?q="+query)
time.sleep(5)

data = []       # Will store all the urls after finding

Soup = BeautifulSoup(browser.page_source, "html.parser")

raw_text = []
remove_urls = []

mydoc = Document()


# Will Search urls for the topic in the browser and take the text part of each page
def scrape():

    pageToSearch = len(
        Soup.find_all("a", attrs={"target", "_blank"}))

    for div_f_tag in Soup.find_all("div", attrs={"class", "b_results"}):
        div_S_tags = div_f_tag.find_all("a", attrs={"target", "_blank"})
        print(div_S_tags)
        print(div_f_tag)

        temp_list = []
        for index, a_tag in enumerate(div_S_tags):
            if index == 0:
                temp_list.append(a_tag.find_all("a", href=True)[0]["href"])

            else:
                try:
                    temp_list.append(a_tag.contents[0])

                except:
                    temp_list.append('')

        data.append(temp_list)

    os.system('CLS')
    print("This Process require time. Please be Patient!!!")
    print('\n\nPages To Search: ', pageToSearch)

    try:
        for i in range(len(data)):
            browser = requests.get(data[i][0])

            print('\nPage no: ', i+1)

            p = len(BeautifulSoup(browser.content, "html.parser").find_all("p"))
            print('p: ', p)

            if p != 0:
                time.sleep(2)

                raw_text.append(BeautifulSoup(
                    browser.content, "html.parser").find_all("p"))

                raw_text.append(
                    "                                                    ")
            else:
                time.sleep(2)
                print(f"\nUnable to fetch data from the Url: {str(data[i][0])}")
                print(
                    "Data From this site will not be Recorded. skipping to the next...")
                remove_urls.append(data[i])

    except:
        print('\n\nError Occured!!!!\n\n Unable to fetch result right now.')
        print("Data Which was collected is been Stored\n\n")

    for i in remove_urls:
        idx = data.index(i)
        data.pop(idx)

    if os.path.exists("./Data") == False:
        os.mkdir("Data")

    with open("Data/"+file_name+"_url.csv", "w+") as f:
        csv_Writter = csv.writer(f)
        csv_Writter.writerows(data)


def remove_tags(text):

    soup = BeautifulSoup(text, "html.parser")

    a_tag_words = []
    temp_text = []

    for data in soup(['img']):
        data.decompose()
        temp_text.append(' '.join(soup.stripped_strings))

    for data in soup("a"):
        a_tag_words.append(data)
        temp_text.append(' '.join(soup.stripped_strings))

    # print('\n\n\n', a_tag_words)

    # temp_text.append([i.text.rstrip() for i in soup.find_all("p")])

    para = temp_text

    try:
        mydoc.add_heading(query)
        mydoc.add_paragraph(para[0])
        mydoc.save("Data/"+file_name+" Data.docx")

    except:
        mydoc.add_heading(query)
        mydoc.add_paragraph(para)
        mydoc.save("Data/"+file_name+" Data.docx")

    print("\n\nData has Been Stored")


scrape()
remove_tags(str(raw_text))
print("\n\n\n\n\n\n")
browser.close()
input('exit...')
exit()
