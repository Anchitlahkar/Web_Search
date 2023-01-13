from selenium import webdriver
from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from webdriver_manager.chrome import ChromeDriverManager
from bs4 import BeautifulSoup
import time
import csv
import requests
from docx import Document
import os
from nlp import summerise


def resource_path(relative_path):
    try:
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.dirname(__file__)
    return os.path.join(base_path, relative_path)


query = input('Key Word: ')
file_name = input('File Name: ')


start_url = "https://www.google.com/"
browser = webdriver.Chrome(service=Service(ChromeDriverManager().install()))


time.sleep(5)

browser.get(start_url)

time.sleep(5)

browser.get(start_url+"search?q="+query)
time.sleep(5)

data = []       # Will store all the urls after finding

Soup = BeautifulSoup(browser.page_source, "html.parser")

raw_text = []
remove_urls = []

text_summery = []
final_summery = ""


mydoc = Document()
mydocNLP = Document()
mydocNPLPages = Document()


# Will Search urls for the topic in the browser and take the text part of each page
def scrape():

    pageToSearch = len(
        Soup.find_all("div", attrs={"class", "tF2Cxc"}))

    for div_f_tag in Soup.find_all("div", attrs={"class", "tF2Cxc"}):
        div_S_tags = div_f_tag.find_all("div", attrs={"class", "yuRUbf"})

        temp_list = []
        for index, a_tag in enumerate(div_S_tags):
            if index == 0:
                temp_list.append(a_tag.find_all("a", href=True)[0]["href"])

            else:
                try:
                    temp_list.append(a_tag.contents[0])

                except:
                    temp_list.append('')

        data.append(temp_list)
        print(f"Try::::{len(data)}")

    os.system('CLS')
    print("This Process require time. Please be Patient!!!")
    print('\n\nPages To Search: ', pageToSearch)

    try:
        for i in range(len(data)):
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 6.0; WOW64; rv:24.0) Gecko/20100101 Firefox/24.0'}
            browser = requests.get(data[i][0], headers=headers)

            print('\nPage no: ', i+1)

            p = len(BeautifulSoup(browser.content, "html.parser").find_all("p"))
            print('p: ', p)

            if p != 0:
                time.sleep(2)

                text = summerise(data[i][0], "")
                text_summery.append(text)
                print("summery: done")

                raw_text.append(BeautifulSoup(
                    browser.content, "html.parser").find_all("p"))

                raw_text.append(
                    "                                                    ")
            else:
                time.sleep(2)
                print(
                    f"\nUnable to fetch data from the Url: {str(data[i][0])}")
                print(
                    "Data From this site will not be Recorded. skipping to the next...")
                remove_urls.append(data[i])

    except ConnectionError as e:
        print("\n\n"+e)
        print("\n\nConnection Error\n\n")
        pass

    # except:
        # print('\n\nError Occured!!!!\n\n Unable to fetch result right now.')
        # print("Data Which was collected is been Stored\n\n")
        # pass

    for i in remove_urls:
        idx = data.index(i)
        data.pop(idx)

    if os.path.exists("./Data") == False:
        os.mkdir("Data")

    if os.path.exists(f"./Data/{file_name}") == False:
        os.makedirs(os.path.join('Data', file_name))

    with open("Data/"+file_name+"/"+file_name+"_url.csv", "w+") as f:
        csv_Writter = csv.writer(f)
        csv_Writter.writerows(data)


def creating_summery_file():
    print("\n\nWriting Summery...")

    raw_summery = ""

    url=0
    for i in text_summery:
        raw_summery += f"  {i}"

        mydocNPLPages.add_heading(data[url][0])
        mydocNPLPages.add_paragraph(raw_summery)
        mydocNPLPages.save("Data/"+file_name+"/"+file_name+"_NLP_Pages.docx")
        
        url+=1

    final_summery = summerise("", raw_summery)

    mydocNLP.add_heading(file_name)
    mydocNLP.add_paragraph(final_summery)
    mydocNLP.save("Data/"+file_name+"/"+file_name+"_NLP.docx")


def remove_tags(text):

    soup = BeautifulSoup(text, "html.parser")

    a_tag_words = []
    temp_text = []

    for data in soup(['img']):
        data.decompose()
        temp_text.append(' '.join(soup.stripped_strings))

    for data in soup("a"):
        a_tag_words.append(data)
        temp_text.append(' '.join(soup.stripped_strings))

    # print('\n\n\n', a_tag_words)

    # temp_text.append([i.text.rstrip() for i in soup.find_all("p")])

    para = temp_text

    creating_summery_file()

    try:
        mydoc.add_heading(query)
        mydoc.add_paragraph(para[0])
        mydoc.save("Data/"+file_name+"/"+file_name+" Data.docx")

    except:
        mydoc.add_heading(query)
        mydoc.add_paragraph(para)
        mydoc.save("Data/"+file_name+"/"+file_name+" Data.docx")

    print("\n\nData has Been Stored")


scrape()
remove_tags(str(raw_text))
print("\n\n\n\n\n\n")
browser.close()
input('exit...')
exit()
